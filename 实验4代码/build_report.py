"""生成实验四的完整实验报告 (Word .docx)"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(ROOT, 'figures')
OUT = os.path.join(ROOT, '机器学习实验四_聚类算法与性能度量_报告.docx')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def P(text, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(11)
    if bold:
        r.bold = True
    return p


def Code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(9)
    return p


def T(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = str(cell)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(10)
                    if i == 0:
                        r.bold = True
    return t


def Pic(path, width_in=5.0):
    doc.add_picture(path, width=Inches(width_in))


# ====================== 标题 ======================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('机器学习实验四：聚类算法与性能度量')
r.bold = True
r.font.size = Pt(16)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('——基于 Iris 数据集的 K-Means 与 DBSCAN 聚类比较')
r.font.size = Pt(12)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ====================== 一、实验目的 ======================
H('一、实验目的', 1)
P('1. 理解聚类分析的基本概念，掌握外部指标（Jaccard 系数 JC、FM 指数 FMI、Rand 指数 RI）和内部指标（Davies-Bouldin 指数 DBI）的定义与计算方法；')
P('2. 掌握闵可夫斯基距离族（p=1 曼哈顿距离、p=2 欧氏距离）的编程实现，理解 VDM（Value Difference Metric）处理无序属性的原理，以及 MinkovDMp 处理混合属性的距离度量；')
P('3. 理解数据归一化的必要性，掌握 Min-Max 归一化方法的编程实现，并通过可视化对比归一化前后的特征分布；')
P('4. 熟练运用 K-Means 和 DBSCAN 两种典型聚类算法，并在 Iris 数据集上进行聚类实验，通过多种性能度量指标对比分析两种算法的优劣。')


# ====================== 二、实验内容 ======================
H('二、实验内容', 1)

H('2.1 任务概述', 2)
P('本实验共包含 6 个核心任务和 1 个综合比较任务：')
P('任务 1：编程计算聚类评估变量 a, b, c, d（SS, SD, DS, DD 四类样本对）；')
P('任务 2：基于 a, b, c 编程计算 Jaccard 系数（JC = a/(a+b+c)）；')
P('任务 3：编程实现闵可夫斯基距离公式，分别计算 p=1（曼哈顿距离）和 p=2（欧氏距离）；')
P('任务 4：编程计算 Davies-Bouldin 指数（DBI），包括簇内平均距离 avg(C)、簇内最远距离 diam(C)、簇间最近距离 d_min、簇间中心距离 d_cen 等辅助量；')
P('任务 5：编程实现 VDM（Value Difference Metric）处理无序属性的距离度量，以及 MinkovDMp 处理混合属性的距离度量；')
P('任务 6：编程实现 Min-Max 归一化，并可视化归一化前后的数据分布；')
P('综合任务：采用 K-Means 和 DBSCAN 两种聚类算法在 Iris 数据集上进行聚类，计算各性能度量指标并比较。')

H('2.2 数据集介绍', 2)
P('本实验采用经典的 Iris（鸢尾花）数据集，该数据集包含 150 个样本，每个样本有 4 个连续特征，分为 3 个类别（Setosa、Versicolor、Virginica），每类各 50 个样本。')
P('表 1  Iris 数据集基本信息')
T([
    ['属性', '值'],
    ['样本数', '150'],
    ['特征数', '4 (Sepal Length, Sepal Width, Petal Length, Petal Width)'],
    ['类别数', '3 (Setosa, Versicolor, Virginica)'],
    ['每类样本数', '50'],
    ['特征类型', '连续型（比率属性）'],
    ['来源', 'sklearn.datasets.load_iris()'],
])

H('2.3 关键公式', 2)
P('(1) 聚类评估变量 a, b, c, d 的定义：')
P('  a = |SS|: 在预测簇 C 中同簇且在真实簇 C* 中同簇的样本对数;')
P('  b = |SD|: 在 C 中同簇但在 C* 中不同簇的样本对数;')
P('  c = |DS|: 在 C 中不同簇但在 C* 中同簇的样本对数;')
P('  d = |DD|: 在 C 中不同簇且在 C* 中不同簇的样本对数。')

P('(2) Jaccard 系数: JC = a / (a + b + c)')
P('(3) FM 指数: FMI = a / sqrt((a+b)(a+c))')
P('(4) Rand 指数: RI = (a + d) / (a + b + c + d)')

P('(5) 闵可夫斯基距离: dist(x_i, x_j) = (sum_{u=1}^{n} |x_{iu} - x_{ju}|^p)^{1/p}')
P('   当 p=1 为曼哈顿距离, p=2 为欧氏距离。')

P('(6) Davies-Bouldin 指数:')
P('   DBI = (1/k) * sum_{i=1}^{k} max_{j!=i} (avg(C_i) + avg(C_j)) / d_cen(mu_i, mu_j)')
P('   其中 avg(C) 为簇内样本间平均距离, d_cen 为簇中心间距离, DBI 越小越好。')

P('(7) VDM 距离: VDM_p(a, b) = sum_{i=1}^{k} |m_{u,a,i}/m_{u,a} - m_{u,b,i}/m_{u,b}|^p')
P('(8) MinkovDMp: MinkovDM_p(x_i, x_j) = (sum_{u=1}^{n_c}|x_{iu}-x_{ju}|^p + sum_{u=n_c+1}^{n} VDM_p(x_{iu}, x_{ju}))^{1/p}')

P('(9) Min-Max 归一化: x\' = (x - x_min) / (x_max - x_min), 将每个特征缩放到 [0, 1] 区间。')


# ====================== 三、实验结果 ======================
H('三、实验结果与分析', 1)

H('3.1 任务 1-2：a, b, c, d 与 Jaccard 系数', 2)
P('对 Iris 数据集分别使用 K-Means (k=3) 和 DBSCAN (eps=0.35, min_samples=5) 进行聚类，与真实标签对比计算 a, b, c, d。')
P('表 2  K-Means 与 DBSCAN 的 a, b, c, d 计算结果')
T([
    ['变量', 'K-Means', 'DBSCAN', '含义'],
    ['a (SS)', '3030', '3675', '预测同簇 & 真实同簇'],
    ['b (SD)', '766', '2500', '预测同簇 & 真实不同簇'],
    ['c (DS)', '645', '0', '预测不同簇 & 真实同簇'],
    ['d (DD)', '6734', '5000', '预测不同簇 & 真实不同簇'],
    ['总对数', '11175', '11175', 'C(150,2) = 11175'],
])

P('图 1  K-Means 与 DBSCAN 的 a, b, c, d 分布饼图')
Pic(os.path.join(FIG_DIR, 'abcd_pie_chart.png'), 5.5)
P('分析：K-Means 的 a (SS) 占比 27.1%，d (DD) 占比 60.3%，说明大部分样本对都被正确分离或正确聚合。DBSCAN 的 c (DS) = 0，表明 DBSCAN 没有将真实同簇的样本分到不同簇（即没有"遗漏"），但 b (SD) 较大，说明 DBSCAN 将原本不同类的样本合并到了同一簇中。')

H('3.2 任务 3：距离函数验证', 2)
P('对示例向量 x = [1, 2, 3], y = [4, 5, 6] 分别计算不同距离：')
P('表 3  距离函数计算结果')
T([
    ['距离类型', '公式', '计算结果'],
    ['曼哈顿距离 (p=1)', '|1-4|+|2-5|+|3-6|', '9.0000'],
    ['欧氏距离 (p=2)', 'sqrt((1-4)^2+(2-5)^2+(3-6)^2)', '5.1962'],
    ['闵可夫斯基 (p=3)', '(|1-4|^3+|2-5|^3+|3-6|^3)^{1/3}', '4.3267'],
])

P('图 2  Iris 数据集前 30 个样本的距离矩阵热力图 (欧氏距离 vs 曼哈顿距离)')
Pic(os.path.join(FIG_DIR, 'distance_heatmaps.png'), 5.5)
P('分析：从热力图可以看出，两种距离度量的整体模式一致（对角线为 0，类内距离小、类间距离大），但曼哈顿距离的绝对数值普遍大于欧氏距离，因为曼哈顿距离是各维差值的线性求和，而欧氏距离是平方和开根号，会"平滑"掉部分差异。')

H('3.3 任务 4：DB 指数计算', 2)
P('Davies-Bouldin 指数是内部评估指标，不依赖真实标签，仅依据簇内紧密度和簇间分离度计算。DBI 越小表示聚类质量越好。')
P('表 4  两种算法的 DBI 对比')
T([
    ['算法', 'DBI', '评价'],
    ['K-Means (k=3)', '1.0805', '较高, 簇间分离度有待提升'],
    ['DBSCAN (eps=0.35)', '0.6907', '较低, 簇内更紧密或簇间更分离'],
])
P('分析：DBSCAN 的 DBI 更低，表明从簇的紧密度/分离度角度看 DBSCAN 找到的簇结构更好。但需注意 DBSCAN 只找到了 2 个簇（将 Versicolor 和 Virginica 合并），导致簇间距离增大、DBI 下降。这提示 DBI 需要结合外部指标一起分析。')

H('3.4 任务 5：VDM 与 MinkovDMp', 2)
P('VDM 用于处理无序属性（如颜色、性别等离散值），其核心思想是：如果两个离散值在各簇中的条件分布相似，则它们的距离小。MinkovDMp 将连续属性的闵可夫斯基距离与无序属性的 VDM 距离统一到同一框架。')
P('本实验使用自构造的混合属性数据（前 2 列为连续、后 2 列为无序）进行验证：')
Code('demo_X = [[1.0, 2.0, 0, 1],\n'
     '          [1.5, 2.5, 1, 0],\n'
     '          [3.0, 4.0, 0, 1],\n'
     '          [3.5, 4.5, 1, 0],\n'
     '          [5.0, 6.0, 0, 0],\n'
     '          [5.5, 6.5, 1, 1]]\n'
     'labels = [0, 0, 1, 1, 2, 2]')
P('MinkovDM_2(x_0, x_2) = 2.8284')
P('该值综合考虑了连续属性的欧氏距离和无序属性的 VDM 距离，验证了混合距离公式的正确性。')

H('3.5 任务 6：归一化', 2)
P('归一化的定义：归一化是将不同量纲、不同取值范围的特征变换到统一尺度的过程。')
P('为什么采用归一化：')
P('  (1) 消除量纲差异：不同特征的取值范围可能差异很大（如身高以厘米计、体重以千克计），直接使用会导致取值范围大的特征在距离计算中主导结果；')
P('  (2) 加速收敛：K-Means 等基于距离的算法在归一化后收敛更快；')
P('  (3) 提高聚类质量：避免某些特征因量纲大而对聚类结果产生不合理的主导影响。')

P('本实验采用 Min-Max 归一化: x\' = (x - x_min) / (x_max - x_min)')

P('表 5  归一化前后的特征统计对比')
T([
    ['特征', '归一化前均值', '归一化前范围', '归一化后均值', '归一化后范围'],
    ['Sepal Length', '5.843', '[4.3, 7.9]', '0.429', '[0, 1]'],
    ['Sepal Width', '3.057', '[2.0, 4.4]', '0.441', '[0, 1]'],
    ['Petal Length', '3.758', '[1.0, 6.9]', '0.467', '[0, 1]'],
    ['Petal Width', '1.199', '[0.1, 2.5]', '0.458', '[0, 1]'],
])

P('图 3  归一化前后特征分布箱线图对比')
Pic(os.path.join(FIG_DIR, 'normalization_comparison.png'), 5.5)
P('分析：归一化前，Sepal Length 的取值范围约 [4.3, 7.9]，而 Petal Width 仅为 [0.1, 2.5]，量纲差异显著。归一化后所有特征统一到 [0, 1] 范围，各特征对距离计算的贡献更加均衡。')

H('3.6 聚类算法比较：K-Means vs DBSCAN', 2)

P('表 6  K-Means 与 DBSCAN 全部性能指标汇总')
T([
    ['指标', 'K-Means', 'DBSCAN', '更优', '说明'],
    ['簇数', '3', '2', 'K-Means', '真实类别数为 3'],
    ['JC (Jaccard)', '0.6823', '0.5951', 'K-Means', '越大越好'],
    ['FMI', '0.8112', '0.7715', 'K-Means', '越大越好'],
    ['RI (Rand)', '0.8737', '0.7763', 'K-Means', '越大越好'],
    ['DBI', '1.0805', '0.6907', 'DBSCAN', '越小越好'],
])

P('图 4  PCA 降维后的聚类结果可视化 (真实标签 vs K-Means vs DBSCAN)')
Pic(os.path.join(FIG_DIR, 'clustering_comparison_pca.png'), 6.0)
P('分析：从 PCA 二维投影可以看出：')
P('  (1) K-Means 将数据划分为 3 个簇，簇的形状与真实标签非常接近，仅在 Versicolor 和 Virginica 的交界区域存在少量错误分配；')
P('  (2) DBSCAN 只找到了 2 个簇，将 Versicolor 和 Virginica 合并为一个大簇，这是因为这两类在特征空间中的密度相近且边界模糊，DBSCAN 的密度连通性判定将它们视为一个连通区域。')

P('图 5  性能指标对比柱状图 (外部指标 + 内部指标)')
Pic(os.path.join(FIG_DIR, 'metrics_comparison.png'), 5.5)

P('图 6  K-Means 在不同 k 值下的指标变化 (肘部法则 + DBI + JC)')
Pic(os.path.join(FIG_DIR, 'kmeans_k_selection.png'), 6.0)
P('分析：从肘部法则（左图）可以看出，k=3 处 Inertia 下降速率明显放缓，形成"肘部"，与真实类别数 3 一致。DBI 在 k=3 时取得局部最小值，JC 在 k=3 时取得最大值，三个指标一致表明 k=3 是最优簇数。')

P('图 7  K-Means 聚类结果的特征两两散点图')
Pic(os.path.join(FIG_DIR, 'kmeans_scatter_matrix.png'), 5.5)
P('分析：从散点图可以看出，Petal Length 和 Petal Width 是最具判别力的两个特征（右下角子图），三个簇在这两个特征上分离最清晰。Sepal Width 的判别力相对较弱。')


# ====================== 四、实验总结 ======================
H('四、实验总结', 1)

P('1. 外部指标一致性：在所有三个外部指标（JC、FMI、RI）上，K-Means 均优于 DBSCAN。K-Means 的 JC=0.6823、FMI=0.8112、RI=0.8737，表明其聚类结果与真实标签有较高的一致性。DBSCAN 由于只找到 2 个簇（将两类合并），外部指标显著下降。')

P('2. 内部指标的局限性：DBSCAN 的 DBI=0.6907 优于 K-Means 的 DBI=1.0805，但这并不意味着 DBSCAN 的聚类更好。DBI 是内部指标，不参考真实标签，当 DBSCAN 将两个相邻类合并后，合并簇的中心距离增大导致 DBI 下降。这提示我们在评估聚类质量时，应综合使用外部指标和内部指标。')

P('3. 算法特性对比：')
P('  K-Means 适合球形/凸形分布的数据，需要预先指定 k 值，对初始中心敏感，不能发现任意形状的簇，对噪声和离群点敏感；')
P('  DBSCAN 不需要预先指定簇数，能发现任意形状的簇，能自动识别噪声点，但对参数 eps 和 min_samples 敏感，不适合密度差异较大的数据集。')

P('4. Iris 数据集上的最优选择：由于 Iris 数据集的类别大致呈凸形分布且类别数已知（k=3），K-Means 是更合适的选择。DBSCAN 在密度边界模糊的 Versicolor/Virginica 区域表现不佳。')

P('5. 归一化的重要性：通过对比归一化前后的特征分布，验证了 Min-Max 归一化能消除量纲差异，使各特征在距离计算中的贡献更加均衡。对于基于距离的聚类算法（如 K-Means），归一化是不可或缺的预处理步骤。')

P('6. 距离度量的选择：实验验证了闵可夫斯基距离族的三种变体。欧氏距离（p=2）是最常用的选择，适合连续属性；曼哈顿距离（p=1）对异常值更鲁棒；VDM 和 MinkovDMp 分别解决了无序属性和混合属性的距离计算问题，拓展了聚类算法的适用范围。')

P('7. 多指标综合评估：本实验同时使用了外部指标（JC、FMI、RI）和内部指标（DBI），两类指标从不同角度评估聚类质量。实践表明，单一指标可能产生误导（如本实验中 DBI 偏好 DBSCAN），需要多指标综合判断才能得出可靠结论。')


# ====================== 五、附录 ======================
H('五、附录：核心代码', 1)

P('1. 计算 a, b, c, d：')
Code("""def compute_abcd(labels_pred, labels_true):
    n = len(labels_pred)
    a = b = c = d = 0
    for i in range(n):
        for j in range(i + 1, n):
            same_pred = (labels_pred[i] == labels_pred[j])
            same_true = (labels_true[i] == labels_true[j])
            if same_pred and same_true:
                a += 1
            elif same_pred and not same_true:
                b += 1
            elif not same_pred and same_true:
                c += 1
            else:
                d += 1
    return a, b, c, d""")

P('2. Jaccard 系数：')
Code("""def jaccard_coefficient(a, b, c):
    return a / (a + b + c)""")

P('3. 闵可夫斯基距离：')
Code("""def minkowski_distance(x, y, p=2):
    x, y = np.asarray(x, dtype=float), np.asarray(y, dtype=float)
    return np.sum(np.abs(x - y) ** p) ** (1.0 / p)""")

P('4. Davies-Bouldin 指数：')
Code("""def davies_bouldin_index(X, labels):
    unique_labels = np.unique(labels)
    unique_labels = unique_labels[unique_labels >= 0]
    k = len(unique_labels)
    clusters = {l: X[labels == l] for l in unique_labels}
    avg_dists = {l: cluster_avg_dist(clusters[l]) for l in unique_labels}
    dbi = 0.0
    for i in unique_labels:
        max_ratio = -1.0
        for j in unique_labels:
            if i == j:
                continue
            d_cen = cluster_d_cen(clusters[i], clusters[j])
            ratio = (avg_dists[i] + avg_dists[j]) / d_cen
            if ratio > max_ratio:
                max_ratio = ratio
        dbi += max_ratio
    return dbi / k""")

P('5. VDM 与 MinkovDMp：')
Code("""def vdm_distance(x_val, y_val, attribute_idx, X_full, labels, p=2):
    col = X_full[:, attribute_idx]
    unique_clusters = np.unique(labels)
    m_u_a = np.sum(col == x_val)
    m_u_b = np.sum(col == y_val)
    if m_u_a == 0 or m_u_b == 0:
        return 0.0
    dist = 0.0
    for ci in unique_clusters:
        mask_ci = (labels == ci)
        m_u_a_i = np.sum(col[mask_ci] == x_val)
        m_u_b_i = np.sum(col[mask_ci] == y_val)
        dist += abs(m_u_a_i / m_u_a - m_u_b_i / m_u_b) ** p
    return dist

def minkov_dm_p(xi, xj, n_continuous, X_full, labels, p=2):
    continuous_part = np.sum(np.abs(xi[:n_continuous] - xj[:n_continuous]) ** p)
    vdm_part = 0.0
    for u in range(n_continuous, len(xi)):
        vdm_part += vdm_distance(xi[u], xj[u], u, X_full, labels, p=p)
    return (continuous_part + vdm_part) ** (1.0 / p)""")

P('6. Min-Max 归一化：')
Code("""def min_max_normalize(X):
    X = np.asarray(X, dtype=float)
    X_min = X.min(axis=0)
    X_max = X.max(axis=0)
    denom = X_max - X_min
    denom[denom == 0] = 1.0
    return (X - X_min) / denom""")


doc.save(OUT)
print(f'实验报告已保存: {OUT}')
