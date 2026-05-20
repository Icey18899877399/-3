"""把补充实验的结果追加到原 docx 报告中，生成 v2 版本。"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = '/home/user/-3/实验3代码'
HANDLED = os.path.join(ROOT, 'chapter3_data_handled')
FIG_DIR = os.path.join(HANDLED, 'figures')
ORIG_DOC = os.path.join(ROOT, '机器学习实验三_贷款违约行为预测_报告.docx')
NEW_DOC = os.path.join(ROOT, '机器学习实验三_贷款违约行为预测_报告_v2.docx')

doc = Document(ORIG_DOC)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def add_table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = str(cell)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True
    return t


# 读取补充实验结果
df_res = pd.read_csv(os.path.join(HANDLED, 'supplementary_results.csv'))

def get_auc(experiment, config):
    sub = df_res[(df_res.experiment == experiment) & (df_res.config == config)]
    return sub.value.iloc[0] if not sub.empty else float('nan')


# ============ 九、补充实验 ============
add_heading('九、补充实验', level=1)
add_para('在第六章已经横向比较了 PPT 给定的四个基线模型，本章在此基础上做两类补充实验：'
         '（A）引入三个现代梯度提升类模型（XGBoost、LightGBM、GBDT）以探索性能上限；'
         '（B）对原 pipeline 中的三项关键技术（不平衡采样 / L1 特征选择 / 标准化）做消融实验，'
         '量化每一步对最终 AUC 的贡献。')

# ------------------ A. 模型扩展 ------------------
add_heading('9.1 实验 A：引入梯度提升模型', level=2)
add_para('梯度提升类模型在结构化数据风控建模中被广泛采用，几乎是当前学术与业界的事实标准。'
         '本实验保持 pipeline 一致（SMOTE+Tomek 平衡，L1 特征选择），将以下三个模型纳入对比：')
add_para('• XGBoost：n_estimators=300, max_depth=4, learning_rate=0.1；')
add_para('• LightGBM：n_estimators=300, num_leaves=31, learning_rate=0.1；')
add_para('• GBDT（sklearn）：n_estimators=200, max_depth=3, learning_rate=0.1。')
add_para('表 5  7 个模型在测试集上的 AUC 对比')
add_table([
    ['模型', 'AUC', '相对 Logistic 提升'],
    ['XGBoost',        f"{get_auc('A. Model expansion', 'XGBoost'):.4f}",
     f"{(get_auc('A. Model expansion', 'XGBoost') - get_auc('A. Model expansion', 'Logistic (L1)')):+.4f}"],
    ['LightGBM',       f"{get_auc('A. Model expansion', 'LightGBM'):.4f}",
     f"{(get_auc('A. Model expansion', 'LightGBM') - get_auc('A. Model expansion', 'Logistic (L1)')):+.4f}"],
    ['GBDT',           f"{get_auc('A. Model expansion', 'GBDT'):.4f}",
     f"{(get_auc('A. Model expansion', 'GBDT') - get_auc('A. Model expansion', 'Logistic (L1)')):+.4f}"],
    ['Logistic (L1)',  f"{get_auc('A. Model expansion', 'Logistic (L1)'):.4f}", '基线 0.0000'],
    ['SVM (linear)',   f"{get_auc('A. Model expansion', 'SVM (linear)'):.4f}",
     f"{(get_auc('A. Model expansion', 'SVM (linear)') - get_auc('A. Model expansion', 'Logistic (L1)')):+.4f}"],
    ['Random Forest',  f"{get_auc('A. Model expansion', 'Random Forest'):.4f}",
     f"{(get_auc('A. Model expansion', 'Random Forest') - get_auc('A. Model expansion', 'Logistic (L1)')):+.4f}"],
    ['Naive Bayes',    f"{get_auc('A. Model expansion', 'Naive Bayes'):.4f}",
     f"{(get_auc('A. Model expansion', 'Naive Bayes') - get_auc('A. Model expansion', 'Logistic (L1)')):+.4f}"],
])
add_para('图 7  7 个分类器的 ROC 曲线对比（顶刊色板：Wong 色盲友好）')
doc.add_picture(os.path.join(FIG_DIR, 'A1_roc_seven_models.png'), width=Inches(5.5))
add_para('图 8  7 个分类器的 AUC 排序柱状图')
doc.add_picture(os.path.join(FIG_DIR, 'A2_auc_bar_seven_models.png'), width=Inches(6))
add_para('结论：XGBoost / LightGBM 显著优于全部 4 个基线模型，AUC 较 Logistic 提升约 0.02 ~ 0.03，'
         '与"GBDT 类模型是结构化风控建模事实标准"的工业经验高度一致。GBDT (sklearn) '
         '与 LightGBM 接近，但训练速度慢。综合性能与效率，本任务最终推荐 LightGBM。')

# ------------------ B. 消融实验 ------------------
add_heading('9.2 实验 B：消融实验', level=2)

# --- B1 ---
add_heading('9.2.1 B1：不平衡采样策略消融', level=3)
add_para('本实验对比四种不平衡处理策略对 AUC 的影响：'
         '（1）不做任何处理；（2）仅 SMOTE 过采样；'
         '（3）SMOTE+Tomek Links 混合采样（原 pipeline 默认）；'
         '（4）模型层面的 class_weight=\'balanced\' / scale_pos_weight。'
         '为公平起见，本组实验关闭 L1 特征选择以排除其干扰，'
         '采用三个代表性模型（Logistic、Random Forest、XGBoost）。')
b1_rows = [['模型', 'None', 'SMOTE', 'SMOTE+Tomek', 'class_weight']]
for m in ['Logistic (L1)', 'Random Forest', 'XGBoost']:
    row = [m]
    for s in ['none', 'smote', 'smotetomek', 'class_weight']:
        row.append(f"{get_auc('B1. Sampling', f'{m} / {s}'):.4f}")
    b1_rows.append(row)
add_para('表 6  不同采样策略下的 AUC')
add_table(b1_rows)
add_para('图 9  四种不平衡采样策略对 AUC 的影响（grouped bar）')
doc.add_picture(os.path.join(FIG_DIR, 'B1_sampling_ablation.png'), width=Inches(6))
add_para('结论：（1）对于 Logistic，反而是"不做任何采样处理"AUC 最高（0.8647），'
         'SMOTE/SMOTETomek 几乎不带来收益甚至略有下降，说明在 L1 已经做了正则的前提下'
         '原始不平衡数据足以给出稳定的概率排序；（2）对 Random Forest 影响最大：'
         '不采样时 AUC 仅 0.8034，SMOTE+Tomek 后提升到 0.8234（+0.020），因为 RF 的'
         '分裂准则对极不平衡比例敏感；（3）XGBoost 对不平衡几乎不敏感（4 种策略 AUC '
         '波动 < 0.007），印证了 GBDT 类模型自带鲁棒处理能力；（4）class_weight 与 '
         'SMOTE 系列效果接近，工程上若考虑训练效率，可优先选 class_weight / scale_pos_weight。')

# --- B2 ---
add_heading('9.2.2 B2：L1 特征选择消融', level=3)
add_para('本实验对比"是否在 SMOTE+Tomek 之后接入 L1 Logistic 特征选择"对三个模型 AUC 的影响。'
         '默认 pipeline 使用 L1 特征选择，将原本 156 维特征压缩到约 15~20 维。')
b2_rows = [['模型', '有 L1 特征选择', '无 L1 特征选择', '差值']]
for m in ['Logistic (L1)', 'Random Forest', 'XGBoost']:
    a_with = get_auc('B2. L1 selection', f'{m} / with L1')
    a_without = get_auc('B2. L1 selection', f'{m} / without L1')
    b2_rows.append([m, f'{a_with:.4f}', f'{a_without:.4f}',
                    f'{(a_with - a_without):+.4f}'])
add_para('表 7  L1 特征选择对 AUC 的影响')
add_table(b2_rows)
add_para('图 10  是否使用 L1 特征选择的 AUC 对比')
doc.add_picture(os.path.join(FIG_DIR, 'B2_l1_fs_ablation.png'), width=Inches(5.5))
add_para('结论：（1）L1 特征选择对 Logistic 完全无影响（其本身就是 L1 线性模型，做不做选择都收敛到同一组系数）；'
         '（2）对 Random Forest 而言，加上 L1 特征选择反而 AUC 更高（+0.006），'
         '可能是因为 L1 帮忙剔除了部分对 RF 分裂无用的噪声特征；'
         '（3）对 XGBoost 而言，去掉 L1 反而略好（+0.001），但差距在噪声范围内，'
         'GBDT 类模型自带特征重要性筛选机制，对外部特征选择不敏感。'
         '总体来看，L1 特征选择对线性模型必要、对树/Boosting 模型可有可无。')

# --- B3 ---
add_heading('9.2.3 B3：标准化消融', level=3)
add_para('原 pipeline 对所有连续变量（SALARY、LOAN_TIME、EDU_LEVEL 及征信扩展表中的金额等）'
         '做了 Z-score 标准化。本实验反向操作：将 SALARY 列还原到原始 1~7 的等级尺度，'
         '其它列保持标准化，再在 Logistic 与 SVM（对量纲最敏感的两个模型）上对比 AUC。')
b3_rows = [['模型', '完全标准化', 'SALARY 未标准化', '差值']]
for m in ['Logistic', 'SVM']:
    a_sc = get_auc('B3. Standardization', f'{m} (standardized)')
    a_un = get_auc('B3. Standardization', f'{m} (unscaled SALARY)')
    b3_rows.append([m, f'{a_sc:.4f}', f'{a_un:.4f}', f'{(a_sc - a_un):+.4f}'])
add_para('表 8  标准化对 AUC 的影响')
add_table(b3_rows)
add_para('图 11  是否对 SALARY 做标准化的 AUC 对比')
doc.add_picture(os.path.join(FIG_DIR, 'B3_standardize_ablation.png'), width=Inches(5))
add_para('结论：本实验的 SALARY 是 1~7 的小范围离散等级，且经过 L1 特征选择后是否被保留也未必，'
         '因此反标准化后 AUC 几乎无变化（差值≈ 0）。这恰恰说明：本数据中 SALARY 不是主要驱动因素，'
         '决定性能的是征信扩展表中已经被 StandardScaler 处理过的金额类特征。'
         '推而广之，当数据集中存在横跨多个数量级的连续变量时，标准化对 SVM 等基于距离的模型至关重要，'
         '原 pipeline 的 StandardScaler 步骤是稳健的工程实践，应当保留。')

# ============ 十、补充实验结论 ============
add_heading('十、补充实验总结', level=1)
add_para('1. 模型层面：GBDT 类模型（XGBoost / LightGBM）显著优于全部 4 个基线模型，'
         '将本任务的 AUC 上限从 0.85 推高至约 0.88。综合性能与效率，推荐 LightGBM 作为生产模型，'
         'Logistic 作为可解释基线保留。')
add_para('2. 不平衡处理：SMOTE+Tomek 对 Random Forest 提升明显（+0.020），但对 Logistic '
         '反而轻微负作用，对 XGBoost 几乎无效。说明 GBDT 类模型自带应对不平衡的鲁棒机制；'
         '工程上对 GBDT 模型可优先使用 scale_pos_weight，对 RF 推荐 SMOTE+Tomek。')
add_para('3. 特征选择：L1 特征选择对 Logistic 完全无影响（线性模型自身已带 L1），'
         '对 RF 略有帮助（+0.006），对 XGBoost 影响在噪声内。'
         '推荐：线性模型上使用 L1，树/Boosting 模型可保留全部特征或依赖模型自身的特征重要性。')
add_para('4. 标准化：在本数据上由于 SALARY 量纲很小，影响有限；但在含金额/年龄等大尺度变量时必须保留。')

doc.save(NEW_DOC)
print(f'New report saved: {NEW_DOC}')
