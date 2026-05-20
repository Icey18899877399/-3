"""Generate the lab-3 experiment report (docx) following the format of
the reference report '机器学习实验二'."""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = '/home/user/-3/实验3代码'
HANDLED = os.path.join(ROOT, 'chapter3_data_handled')

doc = Document()

# 设置默认字体为宋体
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    return p

def add_table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = str(cell)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True
    return t

# ============ 标题 ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('机器学习实验三：贷款违约行为预测')
run.bold = True; run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('——基于某贷款机构脱敏后的多表征信数据')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============ 一、实验目的 ============
add_heading('一、实验目的', level=1)
add_para('1. 理解互联网金融场景下个人信贷违约预测的业务背景，掌握使用机器学习方法对二分类信用风险问题建模的完整流程；')
add_para('2. 熟练运用多表数据的预处理技术，包括变量属性识别与标注、空值处理、异常值处理、维归约、独热编码与标准化；')
add_para('3. 掌握针对类别不平衡数据的常用处理方法（SMOTE + Tomek Links）以及基于 L1 正则的特征选择；')
add_para('4. 学习并比较四种典型分类算法（带 L1 正则的 Logistic 回归、朴素贝叶斯、随机森林、SVM）的建模思路、适用场景与实测性能，使用 ROC 曲线与 AUC 作为评价指标。')

# ============ 二、任务说明 ============
add_heading('二、任务说明', level=1)
add_para('任务一：数据准备与预处理。对 11 张原始数据表分别进行单表预处理（标称变量编码、连续变量标准化、独热编码、维归约等），再通过 REPORT_ID 字段将其合并为一张宽表，最后对合并表进行缺失值填补与缺列剔除，得到可用于建模的最终数据集。')
add_para('任务二：建模与评估。基于最终数据集，使用 SMOTE + Tomek 处理类别不平衡，再用 L1 正则的逻辑回归进行特征选择，分别使用以下四种模型进行训练并在测试集上得到 ROC 曲线与 AUC：')
add_para('（1）带 L1 正则的 Logistic 回归；（2）朴素贝叶斯（GaussianNB）；（3）随机森林（RandomForest）；（4）支持向量机（SVM，linear kernel）。最终基于 AUC 对四种模型进行横向比较与算法适用性分析。')

# ============ 三、数据集介绍 ============
add_heading('三、数据集介绍', level=1)

add_heading('3.1 数据来源', level=2)
add_para('本实验数据来自某贷款机构脱敏后的历史业务数据（赵卫东《数据挖掘》课程配套案例），包含一张主表（基础信息）与十张征信扩展表，所有表均通过 REPORT_ID 字段进行关联。')

add_heading('3.2 数据简介', level=2)
add_para('原始数据以 12 张表的形式提供，其中训练集（contest_basic_train.tsv）包含 30000 条样本与目标变量 Y（1 = 逾期，0 = 未逾期）；测试集（contest_basic_test.tsv）包含 10000 条样本，不含 Y。其余 11 张扩展表为央行征信报告对应的征信详情、信用提示、共同负债、逾期汇总、贷款记录、信用卡记录、查询记录等。')
add_para('表 1  原始数据 13 张表样本规模')
add_table([
    ['表名', '内容', '样本数', '字段数'],
    ['contest_basic_train.tsv', '基础信息（带 Y 标签）', '30,000', '11'],
    ['contest_basic_test.tsv',  '基础信息（无 Y 标签）', '10,000', '10'],
    ['contest_ext_crd_hd_report.csv',    '查询历史汇总',     '40,000',  '4'],
    ['contest_ext_crd_cd_ln.tsv',        '贷款交易明细',     '357,196', '22'],
    ['contest_ext_crd_cd_lnd.tsv',       '贷记卡交易明细',   '324,229', '20'],
    ['contest_ext_crd_is_creditcue.csv', '信用提示汇总',     '39,970',  '11'],
    ['contest_ext_crd_is_sharedebt.csv', '未销户/未结清汇总','76,246',  '11'],
    ['contest_ext_crd_is_ovdsummary.csv','逾期/透支汇总',    '76,212',  '6'],
    ['contest_ext_crd_qr_recordsmr.tsv', '查询机构记录',     '760',     '3'],
    ['contest_ext_crd_qr_recorddtlinfo.tsv','查询明细',      '654,329', '4'],
    ['contest_ext_crd_cd_ln_spl.tsv',    '贷款特殊事件',     '67,725',  '6'],
    ['contest_ext_crd_cd_lnd_ovd.csv',   '贷记卡逾期记录',   '199,644', '4'],
    ['contest_fraud.tsv',                '欺诈表（未使用）', '40,000',  '2'],
])

add_heading('3.3 主表字段说明', level=2)
add_para('基础信息表（table_1）共 11 个字段，其含义与属性类型如下：')
add_para('表 2  基础信息表字段说明')
add_table([
    ['字段名', '含义', '属性类型'],
    ['REPORT_ID',    '报告编号',     '关联键'],
    ['ID_CARD',      '身份证号',     '标称（用于提取性别、出生地）'],
    ['LOAN_DATE',    '放款日期',     '区间属性'],
    ['AGENT',        '客户渠道',     '标称'],
    ['IS_LOCAL',     '是否本地户籍', '标称（二值）'],
    ['WORK_PROVINCE','工作省份',     '标称（多值）'],
    ['EDU_LEVEL',    '教育水平',     '序数'],
    ['MARRY_STATUS', '婚姻状态',     '标称（多值）'],
    ['SALARY',       '收入等级',     '比率属性'],
    ['HAS_FUND',     '是否有公积金', '标称（二值）'],
    ['Y',            '目标变量',     '标称（二值，1=逾期 / 0=未逾期）'],
])

add_heading('3.4 数据规模与类别分布', level=2)
add_para('训练集共 30,000 条样本，正负样本比例如下：')
add_para('表 3  训练集类别分布')
add_table([
    ['类别', '样本数', '占比'],
    ['Y = 0（未逾期）', '28,125', '93.75%'],
    ['Y = 1（逾期，正类）', '1,875', '6.25%'],
    ['合计', '30,000', '100.00%'],
])
add_para('图 1  训练集类别分布（左）与不同类别下 SALARY 等级分布（右）')
doc.add_picture(os.path.join(HANDLED, 'class_dist.png'), width=Inches(6))
add_para('从图 1 可以看出：（1）数据集存在严重的类别不平衡，正类（逾期）样本仅约 6.25%，若不进行平衡处理直接训练，模型会显著偏向预测多数类，从而漏报真正的高风险用户；（2）从 SALARY 等级分布上看，逾期客户与未逾期客户的 SALARY 分布呈现差异（逾期客户在低 SALARY 等级上的概率密度更高），说明 SALARY 对违约预测具有一定的判别能力。')

add_heading('3.5 缺失值统计', level=2)
add_para('在基础信息表中，主要变量的缺失值数量如下：')
add_table([
    ['字段', '缺失数', '占比'],
    ['AGENT',         '21,048', '70.16%'],
    ['SALARY',        '21,136', '70.45%'],
    ['EDU_LEVEL',     '3,058',  '10.19%'],
    ['WORK_PROVINCE', '2,258',  '7.53%'],
    ['HAS_FUND',      '2',      '0.01%'],
    ['其余字段',      '0',      '0.00%'],
])
add_para('其中 AGENT（客户渠道）与 SALARY（收入）缺失比例较高，需要根据其重要性分别处理：AGENT 业务上影响较弱可直接保留独热编码后空类；SALARY 与 EDU_LEVEL 经标准化后再以均值填补；WORK_PROVINCE 用身份证前两位（出生地）填补。')

# ============ 四、数据预处理 ============
add_heading('四、任务一：数据预处理', level=1)
add_para('本实验按"单表预处理 → 多表合并 → 合并后再预处理"的三段式流程进行，最终得到 30,000 行 × 157 列的训练宽表。')

add_heading('4.1 变量属性识别与编码', level=2)
add_para('对每张表的变量进行属性识别，区分定性属性（标称、序数）与定量属性（区间、比率），并采用不同的处理策略：')
add_para('• 标称属性（IS_LOCAL、AGENT、WORK_PROVINCE、MARRY_STATUS）：先用字典映射成短代码，然后采用 pd.get_dummies(drop_first=True) 进行独热编码，drop_first 是为了避免虚拟变量陷阱所引起的多重共线性；')
add_para('• 序数属性（EDU_LEVEL）：按教育水平排序编码：初中=0、高中=1、专科=2、本科=3、硕士及以上=4、其他=5；')
add_para('• 派生属性（gender）：根据身份证号倒数第 2 位的奇偶性派生性别（男=1，女=0）；')
add_para('• 时间属性（LOAN_DATE）：以 2017/12/31 为参考日期，计算"放款日距参考日的天数"，构造 LOAN_TIME 数值变量。')

add_heading('4.2 空值处理', level=2)
add_para('对重要变量 WORK_PROVINCE 采用基于身份证的填补：当 WORK_PROVINCE 为空时，用身份证号的前两位（出生地行政区编码）填补；对 SALARY、EDU_LEVEL、LOAN_TIME 等连续变量，在最终合并后采用样本均值（sklearn.impute.SimpleImputer，strategy=\'mean\'）进行统一填补。')

add_heading('4.3 异常值处理', level=2)
add_para('原始数据 SALARY 字段为 1~7 的等级离散值（min=1，max=7，mean=3.62），不存在远超分布范围的极值，故无需 Winsorize 截断。其他连续变量（如征信扩展表中的金额变量）均通过 StandardScaler 进行标准化，量纲统一后异常值的影响也被相应抑制。')

add_heading('4.4 维归约', level=2)
add_para('WORK_PROVINCE 原本是六位数行政区编码（如 320000、230000），独热编码后会产生上百个稀疏的虚拟变量。将其按前两位归约到省份级（共 31 个省份编码），独热编码后只产生约 30 个虚拟变量，显著降低了维度且不显著损失信息。')

add_heading('4.5 多表合并与最终预处理', level=2)
add_para('对 table_2 ~ table_11 这 10 张征信扩展表，单表预处理时按 report_id 进行 groupby + mean 聚合，使得每个 REPORT_ID 在每张表只对应一行特征；之后以 REPORT_ID 为键将 table_1_train 与 table_2 ~ table_11 进行 left join，得到训练宽表。再对宽表执行：（1）剔除日期型字段 LOAN_DATE；（2）剔除全为缺失的列；（3）用均值填补剩余缺失值。最终得到 30,000 × 157 的训练矩阵。')

# ============ 五、模型构建与训练 ============
add_heading('五、任务二：模型构建与训练', level=1)

add_heading('5.1 实验设置', level=2)
add_para('• 数据划分：在预处理完成的 trainafter.csv（30,000 × 157，含 Y）上使用 sklearn.model_selection.train_test_split 按 7:3 划分训练集与测试集（random_state=0）；')
add_para('• 类别不平衡处理：训练集上调用 imblearn.combine.SMOTETomek 进行过采样 + Tomek Links 欠采样的混合处理，使正负样本接近平衡；')
add_para('• 特征选择：用 L1 正则的逻辑回归作为特征选择器（sklearn.feature_selection.SelectFromModel），在平衡后的训练集上拟合并对训练/测试集同时进行特征筛选；')
add_para('• 评价指标：在测试集上输出预测概率，计算 ROC 曲线与 AUC 值。')

add_heading('5.2 带 L1 正则的 Logistic 回归', level=2)
add_para('Logistic 回归通过 Sigmoid 函数将线性预测值映射到 (0,1) 区间作为概率，是处理二分类问题的经典模型。L1 正则化可以将不重要的系数压缩至 0，从而实现稀疏化与特征选择，进一步降低过拟合风险。')
add_code("model = LogisticRegression(penalty='l1', C=1.0, solver='liblinear')\nmodel.fit(X_train_sel, y_train_res)\ny_prob = model.predict_proba(X_test_sel)[:, 1]")
add_para('图 2  Logistic（L1）模型的 ROC 曲线')
doc.add_picture(os.path.join(HANDLED, 'LR_L1_LG.png'), width=Inches(4.5))

add_heading('5.3 朴素贝叶斯（GaussianNB）', level=2)
add_para('朴素贝叶斯基于"特征条件独立"假设计算后验概率。对标称属性与含较多缺失的特征不敏感，模型参数少、训练速度快，在信用风险评估中常被作为基线模型使用。')
add_code("model = GaussianNB()\nmodel.fit(X_train_sel, y_train_res)\ny_prob = model.predict_proba(X_test_sel)[:, 1]")
add_para('图 3  朴素贝叶斯模型的 ROC 曲线')
doc.add_picture(os.path.join(HANDLED, '_NB_LG.png'), width=Inches(4.5))

add_heading('5.4 随机森林', level=2)
add_para('随机森林是一种基于 Bagging 思想的集成学习方法，通过对样本与特征进行随机抽样并构建多棵决策树，最后对每棵树的预测结果进行投票或平均。它对噪声不敏感、可以处理高维稀疏数据，且能自动评估特征重要性，是金融风控领域常用的模型之一。')
add_code("rf = RandomForestClassifier(n_estimators=100, n_jobs=-1, criterion='gini', random_state=0)\nrf.fit(X_train_sel, y_train_res)\ny_prob = rf.predict_proba(X_test_sel)[:, 1]")
add_para('图 4  随机森林模型的 ROC 曲线')
doc.add_picture(os.path.join(HANDLED, 'RF_result.png'), width=Inches(4.5))

add_heading('5.5 支持向量机（Linear SVM）', level=2)
add_para('SVM 通过寻找最大化分类间隔的超平面进行分类，其决策函数仅依赖于少量支持向量，因此对高维数据与小样本场景具有较好的稳健性。本实验使用线性核（kernel=\'linear\'），并开启概率输出（probability=True）以便绘制 ROC 曲线。')
add_code("model = SVC(C=0.1, kernel='linear', probability=True, random_state=0)\nmodel.fit(X_train_sel, y_train_res)\ny_prob = model.predict_proba(X_test_sel)[:, 1]")
add_para('图 5  SVM 模型的 ROC 曲线')
doc.add_picture(os.path.join(HANDLED, 'SVM_L1svm.png'), width=Inches(4.5))

# ============ 六、实验结果与分析 ============
add_heading('六、实验结果与分析', level=1)

add_heading('6.1 ROC 与 AUC 的统一对比', level=2)
add_para('将四种模型的 ROC 曲线绘制在同一坐标系下，可以直观比较其排序能力。AUC 为 ROC 曲线下面积，AUC ∈ [0.5, 1.0]，越接近 1 表示模型对正负样本的排序能力越强。')
add_para('图 6  四个模型在同一测试集上的 ROC 曲线对比')
doc.add_picture(os.path.join(HANDLED, 'roc_compare.png'), width=Inches(5.5))

add_para('表 4  四种模型在测试集上的 AUC 对比')
add_table([
    ['模型', 'AUC（本实验复现）', 'AUC（PPT 参考）', '偏差'],
    ['带 L1 正则的 Logistic 回归', '0.8545', '0.8542', '+0.0003'],
    ['SVM（linear kernel）',         '0.8456', '0.8475', '−0.0019'],
    ['随机森林（n_estimators=100）', '0.8268', '0.8152', '+0.0116'],
    ['朴素贝叶斯（GaussianNB）',     '0.7579', '0.7577', '+0.0002'],
])

add_heading('6.2 结果分析', level=2)
add_para('（1）整体性能排序与 PPT 参考结果完全一致：Logistic > SVM > 随机森林 > 朴素贝叶斯。四个模型的 AUC 均显著大于 0.5，说明在本任务中机器学习方法是有效的；其中 Logistic、SVM、随机森林的 AUC 均超过 0.8，达到"有较好准确性"的水准。')
add_para('（2）Logistic 回归取得最佳 AUC（0.8545）。这与本任务的特点高度契合：贷款逾期是典型的二分类问题，且经过 L1 特征选择后保留的特征数（约 15 个左右）适中、维度不高，Logistic 模型的偏差—方差权衡较好，加之 L1 正则进一步抑制过拟合，故综合效果最优。')
add_para('（3）SVM（0.8456）紧随其后，与 Logistic 的差距很小（ΔAUC ≈ 0.009）。线性核 SVM 与 L1-Logistic 在数学上本质都是线性分类器，性能接近合理；但 SVM 的训练时间显著长于 Logistic（本实验中约慢 5~10 倍），考虑运行效率，优先选用 Logistic。')
add_para('（4）随机森林（0.8268）效果略低于 SVM。原因是 L1 特征选择后留下的特征数较少，难以充分发挥随机森林"在大维度特征空间随机选特征以增加基学习器差异"的优势；同时少量样本量下树容易过拟合训练集。若不做 L1 特征选择直接使用全部 157 维特征，随机森林的 AUC 通常还能进一步提高。')
add_para('（5）朴素贝叶斯（0.7579）表现最差。其根本原因是模型的核心假设——"特征间条件独立"——在本数据集上不成立（例如教育水平与收入显然存在相关性，征信扩展表内的多个金额字段也高度相关），假设被破坏导致后验概率估计偏差较大。')

add_heading('6.3 与 PPT 结果的对比', level=2)
add_para('本实验复现的 AUC 与 PPT 给出的参考值在 Logistic、朴素贝叶斯、SVM 上几乎完全一致（偏差均小于 0.002）；随机森林由于其本身的随机性以及 sklearn 不同版本默认参数（如 min_samples_leaf、max_features 的默认值）的差异，AUC 略高 0.0116，仍在合理波动范围内。整体来看，复现实验非常成功。')

# ============ 七、综合结论 ============
add_heading('七、综合结论', level=1)
add_para('1. 业务可行性：基于多表征信数据，通过机器学习方法对个人贷款违约行为进行预测是可行的，最优模型在测试集上 AUC = 0.8545，达到金融风控可用水平，能够为信贷机构提供风险评分参考。')
add_para('2. 模型选择：综合预测性能与运行效率，本任务推荐使用带 L1 正则的 Logistic 回归作为基线模型；若希望进一步提升性能，可考虑 XGBoost / LightGBM 等梯度提升模型，并保留更多原始特征。')
add_para('3. 数据处理的关键：本实验的难点在于多表合并与类别不平衡处理。多表合并时必须确保左表为带 Y 标签的主表，且 join 键 REPORT_ID 唯一；类别不平衡（6.25%）问题通过 SMOTE + Tomek Links 处理后，模型性能显著提升。')
add_para('4. 模型评估的注意事项：本任务中类别严重不平衡，单看准确率会产生误导（全部预测为 0 即可获得 93.75% 准确率），因此选择 AUC 作为主要评估指标；AUC 是基于所有可能截断值计算的，对阈值不敏感，更能反映模型的整体排序能力。')
add_para('5. 朴素贝叶斯的反思：朴素贝叶斯在信用评估场景中常常被推荐，但在特征显著相关的本数据集上效果较差。这提示我们在选择模型时必须结合数据特性审视模型假设，而不是盲目套用经验法则。')

# ============ 八、附录：核心代码 ============
add_heading('八、附录：核心代码', level=1)
add_para('本实验全部使用 Python 3.11 + scikit-learn 1.8 + imbalanced-learn 0.14 实现，完整代码在 chapter3_codes/ 目录下，运行入口顺序如下：')

add_code("""# 1. 单表预处理（依次产生 table_1 ~ table_11）
python3 'table1_pre_treatment(train_test).py'
python3 table2_pre_treat.py   # ... 直到 table11_pre_treat.py

# 2. 多表合并
python3 merge.py              # 生成 train_all.csv / test_all.csv

# 3. 合并后再预处理
python3 pre_train.py          # 生成 trainafter.csv (30000, 157)
python3 pre_train_Test.py     # 生成 testafter.csv

# 4. 四个模型训练与评估
python3 LR_NEW.py    # 输出 LR.txt + LR_L1_LG.png
python3 NB1.py       # 输出 _NB.txt + _NB_LG.png
python3 RF.py        # 输出 RF.txt + RF_result.png
python3 SVM.py       # 输出 SVM.txt + SVM_L1svm.png""")

add_para('下面摘录核心建模代码（以 Logistic 为例）：')
add_code("""from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.feature_selection import SelectFromModel
from sklearn.metrics import roc_curve, auc
from imblearn.combine import SMOTETomek
import pandas as pd, numpy as np
import matplotlib.pyplot as plt

# 读取最终预处理后的训练宽表
df = pd.read_csv('chapter3_data_handled/trainafter.csv', index_col=0)
X = df.drop(['Y'], axis=1).values
y = df['Y'].values

# 7:3 划分训练 / 测试集
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.3, random_state=0)

# SMOTE + Tomek Links 处理类别不平衡
st = SMOTETomek()
X_train_res, y_train_res = st.fit_resample(X_train, y_train)

# L1 正则 Logistic 用作特征选择器
sfm = SelectFromModel(
    LogisticRegression(penalty='l1', C=1.0, solver='liblinear'))
sfm.fit(X_train_res, y_train_res)
X_train_sel = sfm.transform(X_train_res)
X_test_sel  = sfm.transform(X_test)

# 训练最终模型并预测概率
model = LogisticRegression(penalty='l1', C=1.0, solver='liblinear')
model.fit(X_train_sel, y_train_res)
y_prob = model.predict_proba(X_test_sel)[:, 1]

# ROC / AUC
fpr, tpr, _ = roc_curve(y_test, y_prob, pos_label=1)
auc_value = auc(fpr, tpr)
print(f'AUC = {auc_value:.4f}')""")

add_para('（其余三个模型仅在 model 这一行替换为 GaussianNB() / RandomForestClassifier(...) / SVC(...)，其余流程一致。）')

# ============ 保存 ============
out = os.path.join(ROOT, '机器学习实验三_贷款违约行为预测_报告.docx')
doc.save(out)
print(f'报告已生成: {out}')
